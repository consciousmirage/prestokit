from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.space_before = Pt(0)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p

def add_para(text):
    p = doc.add_paragraph(text)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_category_table(rows_data):
    """rows_data: list of (label, value) tuples"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(rows_data):
        cell_a = table.cell(i, 0)
        cell_b = table.cell(i, 1)
        cell_a.width = Cm(5)
        run_a = cell_a.paragraphs[0].add_run(label)
        run_a.bold = True
        run_a.font.size = Pt(10)
        run_b = cell_b.paragraphs[0].add_run(value)
        run_b.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_shading(cell_a, 'F0F4FF')
            set_cell_shading(cell_b, 'F0F4FF')
    doc.add_paragraph()

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Free Web Tool Market Research', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    run.font.size = Pt(28)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Search Volume, Competition, Monetization & Opportunity Analysis')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('March 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_heading_styled('Executive Summary', 1)
add_para(
    'This report analyzes 10 categories of free web-based tools that people actively search Google for. '
    'The goal: identify where high search demand meets weak existing solutions, creating an opening for a clean, '
    'fast, modern alternative that can win organic traffic and generate revenue through ads, freemium upgrades, or both.'
)
add_para(
    'The research draws on traffic data from Similarweb and Semrush (late 2025/early 2026), publicly reported revenue figures, '
    'competitive landscape analysis, and known keyword volume benchmarks from SEO industry sources.'
)

add_heading_styled('Key Findings at a Glance', 2)
add_bullet('PDF tools are the proven giant: iLovePDF pulls 226M+ monthly visits. The market is massive but competitive.', 'Biggest proven market: ')
add_bullet('Image tools (compress, resize, remove background) have enormous demand. Remove.bg gets 75M+ visits/month. TinyPNG built a business on a single page.', 'Highest growth category: ')
add_bullet('Invoice generators, email signature generators, and QR code generators. High commercial intent, fragmented competition, real monetization potential.', 'Best opportunity for a new entrant: ')
add_bullet('Developer tools (JSON formatter, regex tester) and Pomodoro timers. Low monetization ceiling; users are ad-averse or expect everything free.', 'Avoid unless you have an edge: ')
add_bullet('Build a multi-tool hub (like TinyWow or SmallSEOTools) rather than a single tool. Bundle 5-10 tools in one category to dominate a niche and cross-sell traffic between pages.', 'Recommended strategy: ')

doc.add_page_break()

# ============================================================
# CATEGORY 1: PDF TOOLS
# ============================================================
add_heading_styled('1. PDF Tools (Merge, Split, Convert, Compress)', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"Compress PDF" ~1.5M-2M/mo globally; "Merge PDF" ~1M-1.5M/mo; "PDF to Word" ~2M+/mo; "Split PDF" ~500K/mo'),
    ('Market Size (traffic)', 'iLovePDF: 226M visits/mo; Smallpdf: 36M visits/mo; PDF24: ~15M; Soda PDF, PDF Candy, others'),
    ('Quality of Existing Solutions', 'GOOD. iLovePDF, Smallpdf, and PDF24 are clean, fast, and well-built. Hard to beat on UX.'),
    ('Monetization Potential', 'HIGH. Smallpdf earns $8.3M/year (bootstrapped, 75-person team). iLovePDF earns millions via freemium + ads.'),
    ('Defensibility', 'LOW. Easy to clone the tools. The moat is SEO authority and backlinks (iLovePDF has 10+ years of domain authority).'),
    ('Verdict', 'HARD TO ENTER. Massive market but dominated by strong incumbents. You would need years to rank.'),
])

add_bold_para('The Landscape')
add_para(
    'PDF tools are the gold standard of the "free tool website" model. iLovePDF is arguably the most successful free tool website ever built -- '
    '226 million monthly visits from a Barcelona team of ~63 people. Smallpdf, bootstrapped from Switzerland, hit $8.3M annual revenue serving 500M+ users. '
    'Google searches for "PDF" hit an all-time high in 2024, surpassing even "iPhone" in search interest for the first time in 20 years.'
)
add_para(
    'The problem for a new entrant: the top 3-4 players have accumulated massive domain authority, millions of backlinks, and brand recognition. '
    'Ranking for "compress PDF" or "merge PDF" would take years of sustained effort. Unless you can find an underserved sub-niche (e.g., PDF tools '
    'specifically for a non-English market, or a privacy-focused "no upload" client-side tool), this is not the place to start.'
)

add_bold_para('Monetization Model')
add_bullet('Freemium: free for basic use, paid for batch processing, larger files, desktop app, API access')
add_bullet('Display ads on free tier (typical RPM: $5-15 for this niche)')
add_bullet('Enterprise/team licenses ($5-10/user/month)')

doc.add_page_break()

# ============================================================
# CATEGORY 2: IMAGE TOOLS
# ============================================================
add_heading_styled('2. Image Tools (Compress, Resize, Convert, Remove Background)', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"Remove background" ~2M+/mo; "Compress image" ~800K-1.2M/mo; "Resize image" ~1M+/mo; "Image to PNG/JPG" ~500K+/mo'),
    ('Market Size (traffic)', 'Remove.bg: 75M visits/mo; TinyPNG: 2.8M visits/mo; Canva Image Tools: part of 270M total Canva visits'),
    ('Quality of Existing Solutions', 'MIXED. Remove.bg is excellent. TinyPNG is great but limited to compression. Many generic sites are ugly and ad-heavy.'),
    ('Monetization Potential', 'HIGH. Remove.bg was valuable enough for Canva to acquire. TinyPNG built a sustainable SaaS from a single-page tool.'),
    ('Defensibility', 'MEDIUM. AI-based background removal has a tech moat. Basic compress/resize is easy to clone.'),
    ('Verdict', 'GOOD OPPORTUNITY for a bundled tool. No single site owns "all image tools" the way iLovePDF owns PDF tools.'),
])

add_bold_para('The Landscape')
add_para(
    'This category is fragmented, which is the opportunity. Remove.bg dominates background removal with 75M+ monthly visits (now owned by Canva). '
    'TinyPNG dominates image compression with 2.8M monthly visits and a SaaS model where the free tool is the lead magnet. '
    'But there is no single "iLovePDF for images" -- no one clean hub where you can compress, resize, convert, crop, and remove backgrounds all in one place.'
)
add_para(
    'Several sites try (allimage.tools, tinyimagekit.com, iloveimg.com) but most are either ugly, slow, or ad-riddled. '
    'A clean, fast, modern image tool hub with 5-8 tools could capture significant organic traffic. '
    'The key terms ("remove background," "compress image," "resize image") collectively pull 4-5M+ searches per month globally.'
)

add_bold_para('Why This Is Attractive')
add_bullet('Images make up 40-44% of a webpage\'s weight, so image optimization is an evergreen need')
add_bullet('Every blogger, social media manager, and small business owner needs these tools regularly')
add_bullet('The tools can run entirely client-side (in-browser), meaning minimal server costs')
add_bullet('AI background removal is now accessible via open-source models, closing the tech gap with Remove.bg')

add_bold_para('Monetization Model')
add_bullet('Display ads (RPM $8-20 for this niche due to design/marketing audience)')
add_bullet('Freemium: free for single images, paid for batch processing and higher resolution')
add_bullet('API access for developers ($10-50/month)')
add_bullet('Affiliate links to design tools, hosting, Canva Pro, etc.')

doc.add_page_break()

# ============================================================
# CATEGORY 3: INVOICE / BILLING TOOLS
# ============================================================
add_heading_styled('3. Invoice Generators / Billing Tools', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"Free invoice generator" ~200K-400K/mo; "Invoice template" ~300K-500K/mo; "Invoice maker" ~100K-200K/mo'),
    ('Market Size (traffic)', 'Invoice-generator.com: ~2-3M visits/mo; Zoho Invoice free tool; Wave; QuickBooks free generator; Canva invoices'),
    ('Quality of Existing Solutions', 'MEDIOCRE. Most are either too simple (just a template) or too complex (full accounting software). The sweet spot is underserved.'),
    ('Monetization Potential', 'VERY HIGH. Finance niche has the highest ad RPMs ($20-50+). Users have strong commercial intent.'),
    ('Defensibility', 'LOW-MEDIUM. Easy to build, but sticky if you add features like saved clients, recurring invoices, payment links.'),
    ('Verdict', 'STRONG OPPORTUNITY. High-intent keyword, high RPMs, and the existing solutions have clear UX gaps.'),
])

add_bold_para('The Landscape')
add_para(
    'The invoice generator space sits in a sweet spot: the keyword "free invoice generator" has strong commercial intent '
    '(these are business owners and freelancers -- high-value ad audiences), but the existing solutions are polarized. '
    'On one end, you have dead-simple PDF templates (Canva, Invoice-generator.com) that lack features like saved clients or payment tracking. '
    'On the other end, you have full accounting platforms (QuickBooks, FreshBooks, Wave) that are overkill for someone who just needs to send one invoice.'
)
add_para(
    'The gap: a clean, modern invoice generator that does slightly more than a template (save your business info, auto-number invoices, '
    'calculate tax, look professional) but does not require signing up for a full accounting suite. This is the "Goldilocks zone" that is underserved.'
)

add_bold_para('Why Finance Niche RPMs Are So High')
add_bullet('Advertisers bidding on these users include: QuickBooks, FreshBooks, Stripe, Square, PayPal, accounting firms')
add_bullet('CPC for "invoice software" keywords: $5-15+ per click')
add_bullet('Display ad RPM for finance/business content: $20-50 (vs. $5-10 for general content)')
add_bullet('Affiliate commissions for accounting software: $50-200+ per referral')

doc.add_page_break()

# ============================================================
# CATEGORY 4: TEXT/WRITING TOOLS
# ============================================================
add_heading_styled('4. Text / Writing Tools (Grammar Check, Summarizer, Paraphraser)', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"Grammar checker" ~500K-1M/mo; "Paraphrasing tool" ~1M+/mo; "Summarizer" ~300K-500K/mo; "Readability checker" ~50K-100K/mo'),
    ('Market Size (traffic)', 'Grammarly: 54.6M visits/mo; QuillBot (paraphraser): ~30M+/mo; SmallSEOTools grammar checker: part of 11M total visits'),
    ('Quality of Existing Solutions', 'DOMINATED BY GIANTS. Grammarly is excellent. QuillBot owns paraphrasing. Hard to compete head-on.'),
    ('Monetization Potential', 'HIGH if you can get traffic. Education/writing niche has decent RPMs ($10-25).'),
    ('Defensibility', 'LOW for basic tools. AI APIs make it easy to build, but also easy to copy. Grammarly has massive brand moat.'),
    ('Verdict', 'AVOID head-on competition with Grammarly/QuillBot. Niche sub-tools (readability scorer, word counter, headline analyzer) could work as part of a bundle.'),
])

add_bold_para('The Landscape')
add_para(
    'Grammarly is a $13B company with 54.6M monthly visits. QuillBot (owned by Course Hero) dominates paraphrasing. '
    'These are not markets you want to attack directly. However, there are underserved sub-niches within the writing tools space:'
)
add_bullet('Readability scoring (Flesch-Kincaid, Gunning Fog) -- most free tools are ugly and outdated')
add_bullet('Word/character counters -- simple but high search volume (~200K+/mo)')
add_bullet('Headline analyzers -- CoSchedule has one but it is gated behind email capture')
add_bullet('Text case converters (UPPERCASE, lowercase, Title Case) -- trivial to build, steady search volume')
add_para(
    'The play here is not to build a Grammarly competitor. It is to build a "Swiss Army knife for text" -- '
    'bundle 8-10 small text utilities on one clean site and capture long-tail traffic across all of them.'
)

doc.add_page_break()

# ============================================================
# CATEGORY 5: DEVELOPER TOOLS
# ============================================================
add_heading_styled('5. Developer Tools (JSON Formatter, Regex Tester, Code Formatter)', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"JSON formatter" ~500K-800K/mo; "Regex tester" ~200K-400K/mo; "Base64 decode" ~300K+/mo; "URL encoder" ~200K+/mo'),
    ('Market Size (traffic)', 'jsonformatter.org: 3.5M visits/mo; regex101.com: ~2-3M visits/mo; jsonlint.com, freeformatter.com: 1-3M each'),
    ('Quality of Existing Solutions', 'MIXED. regex101 is excellent. Many JSON/XML formatters are functional but dated-looking.'),
    ('Monetization Potential', 'LOW-MEDIUM. Developers use ad blockers heavily. RPMs are lower ($3-8). Hard to upsell freemium.'),
    ('Defensibility', 'VERY LOW. These tools are trivial to build. The only moat is SEO authority.'),
    ('Verdict', 'SKIP unless you are a developer yourself and want a passion project. Poor revenue-per-visitor ratio.'),
])

add_bold_para('The Landscape')
add_para(
    'Developer tools have respectable traffic -- jsonformatter.org pulls 3.5M monthly visits, regex101 around 2-3M -- but the '
    'monetization is poor. Developers disproportionately use ad blockers (estimated 40-60% of developer audiences block ads), '
    'and they resist paying for simple utilities. regex101 is actually donation-supported and has sponsors, not a traditional revenue model.'
)
add_para(
    'The tools themselves are also trivially easy to build, meaning there is no defensible moat. '
    'A new, prettier JSON formatter might briefly attract attention, but there is no lock-in -- users will switch to whatever Google ranks first. '
    'Unless you plan to build a comprehensive developer toolkit with unique features (like DataZ.tools is attempting), this category '
    'has a poor effort-to-revenue ratio.'
)

doc.add_page_break()

# ============================================================
# CATEGORY 6: BUSINESS TOOLS
# ============================================================
add_heading_styled('6. Business Tools (QR Code Generator, Email Signature Generator, Form Builder)', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"QR code generator" ~2M-3M/mo; "Email signature generator" ~100K-200K/mo; "Form builder" ~200K-400K/mo'),
    ('Market Size (traffic)', 'QR Code Generator sites collectively: 50M+ visits/mo; HubSpot email sig generator: 134K organic visits/mo alone'),
    ('Quality of Existing Solutions', 'QR: saturated with decent options. Email sig: surprisingly bad -- most are ugly or paywalled. Form builder: dominated by Typeform/Google Forms.'),
    ('Monetization Potential', 'HIGH for QR and email sig (business audience = high RPMs). Form builder is crowded.'),
    ('Defensibility', 'QR: low. Email sig: medium (sticky if users save templates). Form builder: low against incumbents.'),
    ('Verdict', 'EMAIL SIGNATURE GENERATOR is the standout opportunity here. QR codes are crowded. Forms are dominated.'),
])

add_bold_para('Why Email Signature Generator Is a Hidden Gem')
add_para(
    'HubSpot\'s free email signature generator alone drives 134,000 organic visits/month and ranks for 5,900 keywords '
    'with 22,700 backlinks. It is one of HubSpot\'s most successful free tools and a core part of their lead generation strategy. '
    'But here is the thing: HubSpot\'s generator is just okay. It is functional but not beautiful, and it is primarily designed to capture leads for HubSpot\'s CRM.'
)
add_para(
    'Most other email signature generators are either: (a) ugly and outdated (htmlsig.com, mail-signatures.com), '
    '(b) freemium with aggressive paywalls (MySignature, Newoldstamp -- basic features locked behind $4-8/mo), or '
    '(c) part of larger platforms that bury the tool. A genuinely free, beautiful, modern email signature generator with '
    '10-15 templates, social icons, and clean HTML output could rank well and attract a high-value business audience.'
)

add_bold_para('QR Code Generator Reality Check')
add_para(
    'QR code usage surged 433% over the past four years, and QR code scans hit 41.77 million in 2025. '
    'The problem: the market is saturated. qr-code-generator.com (now owned by Bitly), ME-QR, QRCodeChimp, Adobe Express, '
    'Canva, Wix -- everyone offers a free QR code generator now. The search volume is massive (~2-3M/mo) but ranking against '
    'these domains is extremely difficult. Unless you have a unique angle (e.g., QR codes with analytics, custom-branded QR codes), skip this.'
)

doc.add_page_break()

# ============================================================
# CATEGORY 7: SEO TOOLS
# ============================================================
add_heading_styled('7. SEO Tools (Meta Tag Generator, Sitemap Generator, Keyword Tool)', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"Meta tag generator" ~100K-200K/mo; "Sitemap generator" ~100K-150K/mo; "Keyword research tool" ~200K+/mo (dominated by branded searches)'),
    ('Market Size (traffic)', 'SmallSEOTools: 11M visits/mo; SEOptimer, XML-sitemaps.com, and others: 500K-2M each'),
    ('Quality of Existing Solutions', 'MOSTLY BAD. SmallSEOTools is functional but ugly and ad-heavy. Many meta tag generators are outdated and clunky.'),
    ('Monetization Potential', 'MEDIUM-HIGH. SEO/marketing audience has good commercial intent. RPM $10-25.'),
    ('Defensibility', 'LOW. Tools are simple to build. SmallSEOTools\' moat is purely SEO authority from years of operation.'),
    ('Verdict', 'MODERATE OPPORTUNITY. A clean alternative to SmallSEOTools could work, especially if bundled with 10+ tools.'),
])

add_bold_para('The Landscape')
add_para(
    'SmallSEOTools is the poster child of this category: 11M monthly visits, built on a portfolio of 100+ simple SEO utilities '
    '(plagiarism checker, grammar checker, word counter, meta tag generator, etc.). The site is functional but heavily ad-loaded '
    'and visually outdated. They boosted organic traffic to 4.9M+ through aggressive content and programmatic SEO strategies.'
)
add_para(
    'The playbook is clear: do not try to compete with Ahrefs, Semrush, or Moz on keyword research. Instead, build the simple, '
    'free utilities that sit at the top of the funnel -- meta tag generators, robots.txt generators, sitemap generators, '
    'Open Graph tag generators, schema markup generators. These are "build once, rank forever" tools with steady search demand.'
)

add_bold_para('Best Sub-Tools to Build')
add_bullet('Meta tag generator (with Open Graph and Twitter Card preview)')
add_bullet('XML sitemap generator')
add_bullet('Robots.txt generator and validator')
add_bullet('Schema markup generator (FAQ, Product, Article)')
add_bullet('Page speed insights / Core Web Vitals checker')
add_bullet('Hreflang tag generator')

doc.add_page_break()

# ============================================================
# CATEGORY 8: SOCIAL MEDIA TOOLS
# ============================================================
add_heading_styled('8. Social Media Tools (Bio Generator, Hashtag Generator, Link-in-Bio)', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"Hashtag generator" ~200K-400K/mo; "Link in bio" ~300K+/mo; "Instagram bio" ~200K+/mo; "Social media scheduler" ~100K-200K/mo'),
    ('Market Size (traffic)', 'Linktree: 1.2B unique visitors/mo; Buffer/Hootsuite hashtag tools: part of larger platforms'),
    ('Quality of Existing Solutions', 'Link-in-bio: Linktree is functional but basic. Hashtag generators: mostly AI-powered now, decent quality. Bio generators: new category, wide open.'),
    ('Monetization Potential', 'MEDIUM. Creator/influencer audience has lower commercial intent than business audience. RPM $5-12.'),
    ('Defensibility', 'Link-in-bio: HIGH (network effects, user lock-in). Hashtag/bio generators: VERY LOW (commodity tools).'),
    ('Verdict', 'LINK-IN-BIO is the winner here but requires significant product investment. Hashtag/bio generators are too commoditized.'),
])

add_bold_para('The Landscape')
add_para(
    'Linktree is the dominant player with $61.6M annual revenue, 50M users, a $1.3B valuation, and 1.2 billion monthly unique visitors. '
    'That traffic figure is staggering -- it makes Linktree one of the most visited websites on the planet. But Linktree has not evolved much '
    'beyond basic link organization, and many creators are looking for alternatives with better customization, monetization features, and analytics.'
)
add_para(
    'Building a Linktree competitor is a real product play, not a "simple free tool" play. It requires user accounts, custom domains, analytics, '
    'integrations, and ongoing maintenance. If you want to go this route, you are building a SaaS product, not a tool website. '
    'The opportunity is real (Linktree alternatives like Beacons, Stan Store, and The Leap are growing) but the effort is substantially higher.'
)
add_para(
    'Hashtag generators and bio generators, on the other hand, are essentially thin wrappers around AI APIs now. Ahrefs, Buffer, Hootsuite, '
    'and dozens of others offer free versions. There is no moat and minimal monetization potential.'
)

doc.add_page_break()

# ============================================================
# CATEGORY 9: PRODUCTIVITY TOOLS
# ============================================================
add_heading_styled('9. Productivity Tools (Pomodoro Timer, Habit Tracker, Kanban Board)', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"Pomodoro timer" ~200K-400K/mo; "Kanban board" ~100K-200K/mo; "Habit tracker" ~100K-200K/mo; "Online notepad" ~200K+/mo'),
    ('Market Size (traffic)', 'Pomofocus.io: ~300K visits/mo; Notion/Trello dominate kanban; no standalone habit tracker dominates'),
    ('Quality of Existing Solutions', 'Pomodoro: Pomofocus is good but very simple. Kanban: impossible to compete with Notion/Trello. Habit tracker: all are mediocre apps.'),
    ('Monetization Potential', 'LOW. Productivity users expect free tools. Ad RPMs are low ($3-8). Freemium conversion is poor.'),
    ('Defensibility', 'VERY LOW. All are trivial to build. No data lock-in for timer/tracker tools.'),
    ('Verdict', 'SKIP. Low traffic ceiling, low monetization, high competition from established products.'),
])

add_bold_para('The Landscape')
add_para(
    'Pomofocus.io is the leading web-based Pomodoro timer with roughly 300,000 monthly visits. It is clean and functional '
    'but very simple. The issue is not that a better Pomodoro timer cannot be built -- it is that the traffic ceiling is low '
    'and users rarely click ads or pay for productivity timers. The same applies to habit trackers and kanban boards: '
    'Notion, Trello, and Asana dominate the space with free tiers that are hard to beat.'
)
add_para(
    'The one exception might be a "minimal online notepad" or "distraction-free writing tool" -- these have steady search volume '
    'and could be useful as part of a larger tool bundle, but they are not strong enough to build a standalone business around.'
)

doc.add_page_break()

# ============================================================
# CATEGORY 10: FINANCE TOOLS
# ============================================================
add_heading_styled('10. Finance Tools (Budget Planner, Investment Calculator, Loan Calculator)', 1)

add_category_table([
    ('Est. Monthly Search Volume', '"Mortgage calculator" ~5M+/mo; "Investment calculator" ~500K+/mo; "Budget calculator" ~200K-300K/mo; "Compound interest calculator" ~300K+/mo'),
    ('Market Size (traffic)', 'Calculator.net: 633K daily users, ~19M/mo; NerdWallet calculators: part of ~50M+ total visits; Bankrate calculators'),
    ('Quality of Existing Solutions', 'GOOD but dated. Calculator.net works but looks like 2010. NerdWallet/Bankrate calculators are clean but buried in content.'),
    ('Monetization Potential', 'EXTREMELY HIGH. Finance is the #1 RPM niche. Calculator.net earns ~$25K+/month from ads alone. CPC for finance keywords: $5-20+.'),
    ('Defensibility', 'LOW-MEDIUM. Calculators are easy to build. NerdWallet and Bankrate have massive domain authority.'),
    ('Verdict', 'HIGH REWARD but HIGH COMPETITION. If you can rank, the RPMs are extraordinary. The challenge is SEO competition from finance media giants.'),
])

add_bold_para('The Landscape')
add_para(
    'Finance calculators are the highest-RPM category in this entire analysis. Calculator.net, a relatively simple site with dozens '
    'of calculators, earns an estimated $25,000+ per month from Google AdSense alone, with 633,000 daily users. '
    'NerdWallet, which uses calculators as content marketing, generates over $500M in annual revenue (though most from affiliate commissions, not ads).'
)
add_para(
    'The opportunity: Calculator.net looks dated. Its UX has not materially changed in over a decade. A modern, beautifully designed '
    'suite of financial calculators (mortgage, investment, budget, retirement, compound interest, loan payoff) could absolutely win traffic -- '
    'but you would be competing against NerdWallet, Bankrate, Investopedia, and other finance media giants for SEO rankings. '
    'The RPMs make this worth attempting, but only if you have a long-term SEO strategy and patience.'
)

add_bold_para('Why Finance RPMs Are So High')
add_bullet('Finance advertisers (banks, brokerages, insurance companies) pay $5-20+ per click')
add_bullet('Display ad RPMs for finance content: $25-50+ (vs. $5-10 for general content)')
add_bullet('Affiliate commissions: credit cards ($50-200/signup), brokerage accounts ($50-100), insurance quotes ($20-50)')
add_bullet('Users have strong commercial intent -- someone using a mortgage calculator is likely about to take out a mortgage')

doc.add_page_break()

# ============================================================
# RANKED OPPORTUNITIES
# ============================================================
add_heading_styled('Ranked Opportunities: Where to Start', 1)

add_para(
    'Based on the intersection of search volume, competition level, existing solution quality, and monetization potential, '
    'here are the categories ranked from best to worst opportunity for a new entrant:'
)

# Create ranking table
table = doc.add_table(rows=11, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ['Rank', 'Category', 'Traffic Potential', 'Competition', 'Revenue Potential']
for i, header in enumerate(headers):
    cell = table.cell(0, i)
    run = cell.paragraphs[0].add_run(header)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '1a1a2e')

data = [
    ['1', 'Invoice Generator + Billing Tools', 'HIGH', 'MEDIUM', 'VERY HIGH ($25-50 RPM)'],
    ['2', 'Image Tool Hub (compress/resize/bg remove)', 'VERY HIGH', 'MEDIUM', 'HIGH ($10-20 RPM)'],
    ['3', 'Email Signature Generator', 'MEDIUM', 'LOW-MED', 'HIGH ($15-30 RPM)'],
    ['4', 'Finance Calculators Bundle', 'VERY HIGH', 'HIGH', 'VERY HIGH ($25-50 RPM)'],
    ['5', 'SEO Mini-Tools Bundle', 'HIGH', 'MEDIUM', 'MEDIUM-HIGH ($10-25 RPM)'],
    ['6', 'PDF Tools Hub', 'VERY HIGH', 'VERY HIGH', 'HIGH ($8-15 RPM)'],
    ['7', 'Link-in-Bio Page Builder', 'VERY HIGH', 'HIGH', 'MEDIUM ($5-12 RPM)'],
    ['8', 'Text Utilities Bundle', 'HIGH', 'HIGH', 'MEDIUM ($8-15 RPM)'],
    ['9', 'Developer Tools Hub', 'MEDIUM', 'MEDIUM', 'LOW ($3-8 RPM)'],
    ['10', 'Productivity Tools', 'LOW', 'HIGH', 'LOW ($3-8 RPM)'],
]

for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        cell = table.cell(row_idx + 1, col_idx)
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if row_idx % 2 == 0:
            set_cell_shading(cell, 'F0F4FF')

doc.add_paragraph()

doc.add_page_break()

# ============================================================
# RECOMMENDED STRATEGY
# ============================================================
add_heading_styled('Recommended Strategy: The Play', 1)

add_heading_styled('Option A: The Invoice + Business Tools Hub (Highest Revenue Per Visitor)', 2)
add_para(
    'Build a clean, modern site with these tools bundled together:'
)
add_bullet('Free invoice generator (the anchor tool -- high search volume, high RPM)')
add_bullet('Free receipt generator')
add_bullet('Free estimate/quote generator')
add_bullet('Free purchase order template')
add_bullet('Email signature generator')
add_bullet('Business card maker')
add_bullet('Profit margin calculator')
add_bullet('Tax calculator')

add_para(
    'Why this works: Every tool targets business owners and freelancers -- the highest-value ad audience. '
    'The RPMs for finance and business tools are 3-5x higher than general content. You are not competing against '
    'one dominant player (unlike PDF tools vs. iLovePDF). The existing solutions are fragmented and mediocre. '
    'And every tool naturally cross-promotes the others (someone generating an invoice also needs a receipt template).'
)

add_heading_styled('Option B: The Image Tools Hub (Highest Traffic Ceiling)', 2)
add_para(
    'Build a clean, modern site with these image tools:'
)
add_bullet('Image compressor (PNG, JPG, WebP, AVIF)')
add_bullet('Image resizer')
add_bullet('Background remover (using open-source AI models)')
add_bullet('Image format converter')
add_bullet('Image cropper')
add_bullet('Image to PDF')
add_bullet('Color picker / palette extractor from image')
add_bullet('Favicon generator')

add_para(
    'Why this works: The combined search volume across all image-related queries is 5M+ per month globally. '
    'The tools can run entirely in the browser (client-side), which means near-zero server costs. '
    'No single site owns this category the way iLovePDF owns PDFs. And the target audience (designers, bloggers, '
    'social media managers) overlaps well with display ad targeting.'
)

add_heading_styled('Option C: Finance Calculator Suite (Highest RPM, Hardest SEO)', 2)
add_para(
    'Only pursue this if you are willing to invest 12-18 months in SEO before seeing significant traffic. '
    'The RPMs are extraordinary ($25-50), but you are competing against NerdWallet, Bankrate, Investopedia, '
    'and Calculator.net -- all of which have massive domain authority. If you can rank, the payoff is huge. '
    'If you cannot rank, you earn nothing. This is the high-risk, high-reward option.'
)

doc.add_page_break()

# ============================================================
# MONETIZATION DEEP DIVE
# ============================================================
add_heading_styled('Monetization Deep Dive', 1)

add_heading_styled('Revenue Model Comparison', 2)

add_bold_para('Display Ads (Easiest to Start)')
add_bullet('Google AdSense: $5-15 RPM for tool sites (varies by niche)')
add_bullet('Ezoic: $10-25 RPM (requires 10K+ monthly pageviews to join)')
add_bullet('Mediavine: $15-40 RPM (requires 50K+ monthly sessions)')
add_bullet('For a site with 100K monthly visitors at $15 RPM = ~$1,500/month')
add_bullet('For a site with 1M monthly visitors at $15 RPM = ~$15,000/month')

add_bold_para('Freemium / Pro Tier (Highest Revenue Per User)')
add_bullet('Free: basic functionality with limits (e.g., 5 invoices/month, single image processing)')
add_bullet('Pro: $5-15/month for unlimited use, batch processing, no ads, saved templates')
add_bullet('Typical conversion rate: 2-5% of free users upgrade')
add_bullet('Smallpdf model: 500M users, $8.3M revenue = very low conversion but massive scale')
add_bullet('For a site with 100K monthly visitors at 3% conversion and $8/mo = $24,000/month (optimistic)')

add_bold_para('Affiliate Marketing (Best for Finance/Business Tools)')
add_bullet('Accounting software affiliates (QuickBooks, FreshBooks, Xero): $50-200 per signup')
add_bullet('Payment processing (Stripe, Square, PayPal): $50-100 per activated account')
add_bullet('Design tool affiliates (Canva Pro, Adobe): $10-50 per signup')
add_bullet('Hosting affiliates (for SEO tools audience): $50-200 per signup')

add_bold_para('Realistic Revenue Projections')

table2 = doc.add_table(rows=5, cols=4)
table2.alignment = WD_TABLE_ALIGNMENT.LEFT
h2 = ['Monthly Visitors', 'Ad Revenue (at $15 RPM)', 'Freemium Revenue (3% at $8/mo)', 'Total Potential']
for i, header in enumerate(h2):
    cell = table2.cell(0, i)
    run = cell.paragraphs[0].add_run(header)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(cell, '1a1a2e')

rev_data = [
    ['10,000', '$150', '$2,400', '$2,550'],
    ['50,000', '$750', '$12,000', '$12,750'],
    ['100,000', '$1,500', '$24,000', '$25,500'],
    ['500,000', '$7,500', '$120,000', '$127,500'],
]
for row_idx, row_data in enumerate(rev_data):
    for col_idx, val in enumerate(row_data):
        cell = table2.cell(row_idx + 1, col_idx)
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(9)
        if row_idx % 2 == 0:
            set_cell_shading(cell, 'F0F4FF')

doc.add_paragraph()
add_para(
    'Note: The freemium revenue numbers assume optimistic conversion rates. In practice, most free tool sites '
    'see 1-3% conversion to paid. Smallpdf, with 500M total users and $8.3M revenue, has an extremely low conversion rate '
    'but makes up for it with massive scale. For a new site, display ads will likely be the primary revenue source '
    'for the first 6-12 months until you build enough traffic and features to justify a paid tier.'
)

doc.add_page_break()

# ============================================================
# TIMELINE
# ============================================================
add_heading_styled('Realistic Timeline & Milestones', 1)

add_bold_para('Month 1-2: Build & Launch')
add_bullet('Build 5-8 tools in your chosen category')
add_bullet('Focus on speed, clean design, mobile-friendly, zero friction (no signups required)')
add_bullet('Set up basic SEO: title tags, meta descriptions, internal linking, sitemap')

add_bold_para('Month 3-6: Content & SEO Foundation')
add_bullet('Create supporting blog content for each tool (how-to guides, templates, examples)')
add_bullet('Target long-tail keywords (e.g., "free invoice template for freelancers" instead of just "invoice generator")')
add_bullet('Build backlinks through guest posts, HARO/Connectively, Reddit, social media')
add_bullet('Expected traffic: 1,000-10,000 monthly visitors')

add_bold_para('Month 6-12: Growth Phase')
add_bullet('Add more tools to the suite (aim for 15-20 total)')
add_bullet('Apply for Ezoic or similar ad network once you hit 10K+ monthly pageviews')
add_bullet('Start capturing emails with lead magnets (templates, guides)')
add_bullet('Expected traffic: 10,000-50,000 monthly visitors')
add_bullet('Expected revenue: $500-2,000/month from ads')

add_bold_para('Month 12-24: Scale & Monetize')
add_bullet('Apply for Mediavine once you hit 50K+ monthly sessions')
add_bullet('Launch freemium/pro tier if usage warrants it')
add_bullet('Expand into adjacent tool categories')
add_bullet('Expected traffic: 50,000-200,000 monthly visitors')
add_bullet('Expected revenue: $2,000-10,000/month')

add_bold_para('Month 24+: Compounding Returns')
add_bullet('SEO compounds: domain authority grows, existing pages rank higher, new pages rank faster')
add_bullet('Websites offering free tools see 35.6% year-over-year growth in organic traffic')
add_bullet('At this stage, you can explore acquisition offers, partnerships, or continued independent growth')

doc.add_page_break()

# ============================================================
# KEY TAKEAWAYS
# ============================================================
add_heading_styled('Final Takeaways', 1)

add_bold_para('1. The model is proven and works.')
add_para(
    'iLovePDF (226M visits/mo), Smallpdf ($8.3M revenue), Calculator.net ($25K+/mo from ads), '
    'Linktree ($61.6M revenue), TinyPNG (millions of users from one page) -- free tool websites '
    'are a legitimate business model with proven revenue paths.'
)

add_bold_para('2. Bundle beats single.')
add_para(
    'Do not build one tool. Build 5-10 tools in a related category. Each tool is a separate page that can rank '
    'on Google, and tools cross-promote each other. SmallSEOTools gets 11M monthly visits with 100+ simple tools.'
)

add_bold_para('3. Design is the differentiator.')
add_para(
    'Most free tool websites look terrible. They are functional but ugly, ad-heavy, and slow. '
    'A clean, fast, modern design with minimal ads is the single biggest competitive advantage you can have. '
    'TinyPNG proved that a single, beautifully designed tool page can build a sustainable business.'
)

add_bold_para('4. Pick the highest-RPM niche you can compete in.')
add_para(
    'Not all traffic is equal. 100K visitors in the finance/business niche generates more revenue than '
    '500K visitors in the developer tools niche. Invoice generators and finance calculators attract users '
    'with strong commercial intent, which means higher ad rates and better affiliate commissions.'
)

add_bold_para('5. SEO is the game. There is no shortcut.')
add_para(
    'Free tool websites live and die by organic search traffic. You need patience (6-12 months to see meaningful traffic), '
    'good on-page SEO, supporting content (blog posts, guides), and backlink building. '
    'The good news: tools naturally attract backlinks because people link to useful resources.'
)

add_bold_para('6. AI has lowered the build barrier but not the distribution barrier.')
add_para(
    'With AI coding tools, you can build a complete tool website in days, not months. '
    'The hard part is no longer building -- it is ranking on Google, which still requires time, authority, and content. '
    'This is actually good news: it means the barrier to entry is not technical ability, but SEO patience and execution.'
)

# ---- SOURCES ----
doc.add_page_break()
add_heading_styled('Sources & Data References', 1)
add_para('Traffic and revenue data sourced from:')
add_bullet('Similarweb traffic analytics (November 2025 - January 2026)')
add_bullet('Semrush website analytics and keyword data (December 2025 - February 2026)')
add_bullet('SiteWorthTraffic revenue estimates')
add_bullet('GetLatka company revenue database (Smallpdf: $8.3M; Linktree: $61.6M)')
add_bullet('Growjo and Tracxn company profiles')
add_bullet('Ahrefs SEO case study on HubSpot free tools strategy')
add_bullet('PDF Association: Google Trends analysis for PDF searches (2024)')
add_bullet('QR Code Tiger: QR code usage statistics (2025)')
add_bullet('Concurate/Buildd: HubSpot SEO analysis (13M monthly organic traffic)')
add_bullet('Monetag, Elementor, MonetizeMore: Website monetization guides (2025-2026)')
add_bullet('Serpzilla, RankTracker: AdSense RPM benchmarks by niche (2025)')
add_bullet('TinyPNG success story via Jolyti.com')
add_bullet('Indie Hackers: TinyWow revenue discussion')
add_bullet('Failory newsletter: Canva traffic analysis (270M visitors/mo)')

# Save
output_path = os.path.expanduser('/Users/nathanfastenberg/Desktop/CLAUDEPROJECTS/Free_Tool_Website_Market_Research.docx')
doc.save(output_path)
print(f"Saved to {output_path}")
